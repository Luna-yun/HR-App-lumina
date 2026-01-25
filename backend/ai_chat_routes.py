from fastapi import APIRouter, HTTPException, status, Depends, UploadFile, File
from motor.motor_asyncio import AsyncIOMotorDatabase
from typing import List, Optional
import logging
import os
import uuid
import hashlib
from datetime import datetime
from qdrant_client import QdrantClient
from qdrant_client.http import models as qmodels
from groq import Groq
import PyPDF2
import docx
import io
from sentence_transformers import SentenceTransformer

from models import (
    KnowledgeDocument, ChatMessage, ChatRequest, ChatResponse
)
from auth_utils import get_current_admin

logger = logging.getLogger(__name__)

router = APIRouter()

# Initialize Groq client for LLM inference
groq_client = Groq(api_key=os.environ.get('GROQ_API_KEY'))

# Initialize Qdrant client
qdrant_client = QdrantClient(
    url=os.environ.get('QDRANT_URL'),
    api_key=os.environ.get('QDRANT_API_KEY'),
)

# Initialize local embedding model (sentence-transformers)
# Using all-MiniLM-L6-v2 - lightweight and fast, 384 dimensions
EMBEDDING_MODEL_NAME = "all-MiniLM-L6-v2"
EMBEDDING_DIM = 384

# Lazy load embedding model to avoid startup delay
_embedding_model = None

def get_embedding_model():
    """Lazy load the embedding model"""
    global _embedding_model
    if _embedding_model is None:
        logger.info(f"Loading embedding model: {EMBEDDING_MODEL_NAME}")
        _embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME)
        logger.info(f"Embedding model loaded successfully. Dimension: {EMBEDDING_DIM}")
    return _embedding_model

COLLECTION_NAME_PREFIX = "hr_knowledge_"

# Dependency to get database
async def get_db() -> AsyncIOMotorDatabase:
    from server import db
    return db


def get_collection_name(company_id: str) -> str:
    """Generate collection name for a company"""
    return f"{COLLECTION_NAME_PREFIX}{company_id.replace('-', '_')}"


def ensure_collection(company_id: str):
    """Ensure Qdrant collection exists for company"""
    collection_name = get_collection_name(company_id)
    try:
        collection_info = qdrant_client.get_collection(collection_name)
        # Check if existing collection has different dimension - recreate if needed
        if collection_info.config.params.vectors.size != EMBEDDING_DIM:
            logger.warning(f"Collection {collection_name} has wrong dimension ({collection_info.config.params.vectors.size} vs {EMBEDDING_DIM}), recreating...")
            qdrant_client.delete_collection(collection_name)
            raise Exception("Need to recreate collection")
    except Exception:
        # Create collection if it doesn't exist or needs recreation
        qdrant_client.create_collection(
            collection_name=collection_name,
            vectors_config=qmodels.VectorParams(
                size=EMBEDDING_DIM,
                distance=qmodels.Distance.COSINE
            )
        )
        logger.info(f"Created Qdrant collection: {collection_name} with dimension {EMBEDDING_DIM}")


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        return text.strip()
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Failed to extract text from PDF"
        )


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text.strip()
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Failed to extract text from DOCX"
        )


def extract_text_from_doc(file_content: bytes) -> str:
    """Extract text from DOC file - limited support"""
    try:
        text = file_content.decode('utf-8', errors='ignore')
        text = ''.join(char for char in text if char.isprintable() or char in '\n\r\t')
        return text.strip()
    except Exception as e:
        logger.error(f"DOC extraction error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Failed to extract text from DOC file. Consider converting to DOCX or PDF."
        )


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """Split text into overlapping chunks"""
    if not text:
        return []
    
    words = text.split()
    chunks = []
    
    for i in range(0, len(words), chunk_size - overlap):
        chunk = ' '.join(words[i:i + chunk_size])
        if chunk:
            chunks.append(chunk)
    
    return chunks


def get_embedding(text: str) -> List[float]:
    """Get embedding using local sentence-transformers model"""
    try:
        model = get_embedding_model()
        
        # Clean and normalize text
        text = text.strip()[:8000]  # Limit text length
        
        # Generate embedding
        embedding = model.encode(text, normalize_embeddings=True)
        
        return embedding.tolist()
        
    except Exception as e:
        logger.error(f"Embedding error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate embedding: {str(e)}"
        )


def get_embeddings_batch(texts: List[str]) -> List[List[float]]:
    """Get embeddings for multiple texts efficiently"""
    try:
        model = get_embedding_model()
        
        # Clean texts
        cleaned_texts = [t.strip()[:8000] for t in texts]
        
        # Generate embeddings in batch (much faster)
        embeddings = model.encode(cleaned_texts, normalize_embeddings=True, show_progress_bar=False)
        
        return embeddings.tolist()
        
    except Exception as e:
        logger.error(f"Batch embedding error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate embeddings: {str(e)}"
        )


@router.post("/admin/chat/upload")
async def upload_knowledge_document(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_admin),
    db: AsyncIOMotorDatabase = Depends(get_db)
):
    """Upload document to knowledge base (Admin only)"""
    try:
        company_id = current_user["company_id"]
        admin_id = current_user["sub"]
        
        # Validate file type
        filename = file.filename.lower()
        if not filename.endswith(('.pdf', '.doc', '.docx')):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Only PDF, DOC, and DOCX files are supported"
            )
        
        # Read file content
        file_content = await file.read()
        file_size = len(file_content)
        
        # Generate content hash
        content_hash = hashlib.md5(file_content).hexdigest()
        
        # Check for duplicate
        existing_doc = await db.knowledge_documents.find_one({
            "company_id": company_id,
            "content_hash": content_hash
        })
        
        if existing_doc:
            return {
                "message": "This document has already been uploaded",
                "document_id": existing_doc["id"],
                "duplicate": True,
                "steps_completed": ["duplicate_check"]
            }
        
        # Extract text based on file type
        if filename.endswith('.pdf'):
            text = extract_text_from_pdf(file_content)
            file_type = "pdf"
        elif filename.endswith('.docx'):
            text = extract_text_from_docx(file_content)
            file_type = "docx"
        else:  # .doc
            text = extract_text_from_doc(file_content)
            file_type = "doc"
        
        if not text or len(text) < 10:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Could not extract meaningful text from document"
            )
        
        # Chunk the text
        chunks = chunk_text(text)
        
        if not chunks:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Document too short or empty"
            )
        
        # Ensure collection exists
        ensure_collection(company_id)
        collection_name = get_collection_name(company_id)
        
        # Create document record
        doc = KnowledgeDocument(
            company_id=company_id,
            filename=file.filename,
            file_type=file_type,
            file_size=file_size,
            content_hash=content_hash,
            chunk_count=len(chunks),
            uploaded_by=admin_id
        )
        
        # Generate embeddings for all chunks (batch processing)
        logger.info(f"Generating embeddings for {len(chunks)} chunks...")
        embeddings = get_embeddings_batch(chunks)
        
        # Store chunks in Qdrant
        points = []
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            point_id = str(uuid.uuid4())
            
            points.append(qmodels.PointStruct(
                id=point_id,
                vector=embedding,
                payload={
                    "document_id": doc.id,
                    "document_name": file.filename,
                    "chunk_index": i,
                    "text": chunk,
                    "company_id": company_id
                }
            ))
        
        # Batch upsert to Qdrant
        qdrant_client.upsert(
            collection_name=collection_name,
            points=points,
            wait=True
        )
        
        # Save document record to MongoDB AFTER successful vector storage
        await db.knowledge_documents.insert_one(doc.dict())
        
        logger.info(f"Uploaded document {file.filename} with {len(chunks)} chunks using local embeddings")
        
        return {
            "message": "Document uploaded and processed successfully",
            "document_id": doc.id,
            "filename": file.filename,
            "chunks_created": len(chunks),
            "duplicate": False,
            "steps_completed": [
                "file_received",
                "text_extracted", 
                "chunks_created",
                "embeddings_generated",
                "vectors_stored",
                "metadata_saved"
            ]
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload document error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"An error occurred while uploading document: {str(e)}"
        )


@router.get("/admin/chat/documents")
async def get_knowledge_documents(
    current_user: dict = Depends(get_current_admin),
    db: AsyncIOMotorDatabase = Depends(get_db)
):
    """Get all knowledge documents for company (Admin only)"""
    try:
        company_id = current_user["company_id"]
        
        documents = await db.knowledge_documents.find({
            "company_id": company_id
        }).sort("created_at", -1).to_list(100)
        
        return {
            "documents": [
                {
                    "id": doc["id"],
                    "filename": doc["filename"],
                    "file_type": doc["file_type"],
                    "file_size": doc["file_size"],
                    "chunk_count": doc["chunk_count"],
                    "created_at": doc["created_at"]
                }
                for doc in documents
            ],
            "total": len(documents)
        }
    
    except Exception as e:
        logger.error(f"Get documents error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="An error occurred"
        )


@router.delete("/admin/chat/documents/{document_id}")
async def delete_knowledge_document(
    document_id: str,
    current_user: dict = Depends(get_current_admin),
    db: AsyncIOMotorDatabase = Depends(get_db)
):
    """Delete knowledge document (Admin only)"""
    try:
        company_id = current_user["company_id"]
        
        doc = await db.knowledge_documents.find_one({"id": document_id})
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )
        
        if doc["company_id"] != company_id:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="You can only delete documents from your company"
            )
        
        # Delete from Qdrant
        collection_name = get_collection_name(company_id)
        try:
            qdrant_client.delete(
                collection_name=collection_name,
                points_selector=qmodels.FilterSelector(
                    filter=qmodels.Filter(
                        must=[
                            qmodels.FieldCondition(
                                key="document_id",
                                match=qmodels.MatchValue(value=document_id)
                            )
                        ]
                    )
                )
            )
            logger.info(f"Deleted document {document_id} from Qdrant")
        except Exception as e:
            logger.warning(f"Failed to delete from Qdrant: {str(e)}")
        
        # Delete from MongoDB
        await db.knowledge_documents.delete_one({"id": document_id})
        
        return {
            "message": "Document deleted successfully",
            "steps_completed": ["vectors_deleted", "metadata_deleted"]
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Delete document error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="An error occurred"
        )


@router.post("/admin/chat", response_model=ChatResponse)
async def chat(
    request: ChatRequest,
    current_user: dict = Depends(get_current_admin),
    db: AsyncIOMotorDatabase = Depends(get_db)
):
    """Chat with AI using RAG from knowledge base (Admin only)"""
    try:
        company_id = current_user["company_id"]
        user_id = current_user["sub"]
        session_id = request.session_id or str(uuid.uuid4())
        
        # Get embedding for query using local model
        query_embedding = get_embedding(request.message)
        
        # Ensure collection exists
        ensure_collection(company_id)
        collection_name = get_collection_name(company_id)
        
        # Search Qdrant for relevant chunks
        search_results = []
        try:
            results = qdrant_client.query_points(
                collection_name=collection_name,
                query=query_embedding,
                limit=5,
                score_threshold=0.3,
                with_payload=True
            )
            search_results = results.points if results and results.points else []
            logger.info(f"Qdrant search returned {len(search_results)} results")
        except Exception as e:
            logger.warning(f"Qdrant search error: {str(e)}")
        
        # Build context from search results
        context_chunks = []
        sources = []
        for result in search_results:
            if result.payload:
                text = result.payload.get("text", "")
                context_chunks.append(text)
                doc_name = result.payload.get("document_name", "Unknown")
                if doc_name not in [s["name"] for s in sources]:
                    sources.append({
                        "name": doc_name,
                        "relevance": round(result.score * 100, 1)
                    })
                logger.info(f"Found relevant chunk from {doc_name} with score {result.score}")
        
        context = "\n\n---\n\n".join(context_chunks) if context_chunks else ""
        
        # Get chat history for this session
        history_messages = await db.chat_messages.find({
            "session_id": session_id,
            "company_id": company_id
        }).sort("created_at", -1).limit(10).to_list(10)
        
        history_messages.reverse()  # Chronological order
        
        # Build conversation history for Groq
        conversation = []
        
        # System prompt
        system_prompt = """You are an intelligent HR assistant with access to company knowledge documents. 
Your role is to help HR administrators with questions about company policies, procedures, and employee-related matters.

CRITICAL INSTRUCTIONS:
1. You MUST base your answers primarily on the provided context from company documents
2. When context is provided, extract and cite specific information from it
3. If the context contains relevant information, summarize and explain it clearly
4. Only if the context truly doesn't contain relevant information, acknowledge this
5. Be specific - quote or paraphrase directly from the documents when possible
6. Always mention which document(s) you're referencing
7. Format your response clearly with bullet points when listing multiple items"""
        
        conversation.append({
            "role": "system",
            "content": system_prompt
        })
        
        # Add history
        for msg in history_messages:
            conversation.append({
                "role": msg["role"],
                "content": msg["content"]
            })
        
        # Add current message with context
        if context:
            user_message = f"""I have retrieved the following relevant information from our company knowledge base:

=== DOCUMENT CONTENT START ===
{context}
=== DOCUMENT CONTENT END ===

Based on the above document content, please answer this question: {request.message}

Important: Your answer should be based on the document content provided above. Quote or reference specific parts of the documents."""
        else:
            user_message = f"""Question: {request.message}

Note: No documents have been uploaded to the knowledge base yet, or no relevant documents were found. Please provide a general response and suggest uploading relevant policy documents for more specific answers."""
        
        conversation.append({
            "role": "user",
            "content": user_message
        })
        
        # Call Groq API
        try:
            chat_completion = groq_client.chat.completions.create(
                messages=conversation,
                model="llama-3.1-8b-instant",
                temperature=0.5,
                max_tokens=2048,
                top_p=1,
            )
            
            assistant_response = chat_completion.choices[0].message.content
        except Exception as e:
            logger.error(f"Groq API error: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Failed to generate AI response"
            )
        
        # Generate reasoning
        if sources:
            reasoning = f"Answer derived from {len(sources)} relevant document(s): {', '.join([s['name'] for s in sources])}. "
            avg_relevance = sum(s['relevance'] for s in sources) / len(sources)
            reasoning += f"Average relevance score: {avg_relevance:.1f}%"
        else:
            reasoning = "No documents in knowledge base matched this query. Response is based on general AI knowledge. Consider uploading relevant policy documents for more specific answers."
        
        # Save user message
        user_chat_message = ChatMessage(
            company_id=company_id,
            user_id=user_id,
            session_id=session_id,
            role="user",
            content=request.message
        )
        await db.chat_messages.insert_one(user_chat_message.dict())
        
        # Save assistant response
        assistant_chat_message = ChatMessage(
            company_id=company_id,
            user_id=user_id,
            session_id=session_id,
            role="assistant",
            content=assistant_response,
            sources=[s["name"] for s in sources]
        )
        await db.chat_messages.insert_one(assistant_chat_message.dict())
        
        return ChatResponse(
            response=assistant_response,
            sources=sources,
            session_id=session_id,
            reasoning=reasoning
        )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Chat error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"An error occurred during chat: {str(e)}"
        )


@router.get("/admin/chat/history")
async def get_chat_history(
    session_id: Optional[str] = None,
    limit: int = 50,
    current_user: dict = Depends(get_current_admin),
    db: AsyncIOMotorDatabase = Depends(get_db)
):
    """Get chat history (Admin only)"""
    try:
        company_id = current_user["company_id"]
        
        query = {"company_id": company_id}
        if session_id:
            query["session_id"] = session_id
        
        messages = await db.chat_messages.find(query).sort("created_at", -1).limit(limit).to_list(limit)
        
        messages.reverse()  # Chronological order
        
        return {
            "messages": [
                {
                    "id": msg["id"],
                    "role": msg["role"],
                    "content": msg["content"],
                    "sources": msg.get("sources", []),
                    "created_at": msg["created_at"]
                }
                for msg in messages
            ],
            "total": len(messages)
        }
    
    except Exception as e:
        logger.error(f"Get chat history error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="An error occurred"
        )


@router.delete("/admin/chat/history")
async def clear_chat_history(
    session_id: Optional[str] = None,
    current_user: dict = Depends(get_current_admin),
    db: AsyncIOMotorDatabase = Depends(get_db)
):
    """Clear chat history (Admin only)"""
    try:
        company_id = current_user["company_id"]
        
        query = {"company_id": company_id}
        if session_id:
            query["session_id"] = session_id
        
        result = await db.chat_messages.delete_many(query)
        
        return {
            "message": f"Deleted {result.deleted_count} messages",
            "deleted_count": result.deleted_count
        }
    
    except Exception as e:
        logger.error(f"Clear chat history error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="An error occurred"
        )
